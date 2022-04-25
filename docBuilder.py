from docx import Document, shared
from pandas import read_excel

data = read_excel("Filenames.xlsx")
documentTitle = "Report_Plots.docx"
imagesOrder = ["position_acceleration", "fx", "fyz", "txz", "ty","three_loads_fx","three_loads_ty"]
document = Document()
font = document.styles["Heading 3"].font
font.size = shared.Pt(12)
font2 = document.styles["Heading 2"].font
font2.size = shared.Pt(13)
font2.underline = False
fontNormal = document.styles["Normal"].font
fontNormal.size = shared.Pt(11)
fontNormal.name = "Calibri"
fontNormal.underline = True
lastFillLevel = 0
# headers = data.loc[0].array
# print(data.shape[0])
iterCounter = 0
caseCounter = 0
print(len(data.columns))
for n in range(1,len(data.columns)):#range(1,2):
	currentFillLevel = int(str(data.columns[n]).split(".")[0])
	currentAcc = data.loc[0].array[n]
	if currentFillLevel != lastFillLevel:
		caseCounter += 1
		if caseCounter <= 2:
			caseStr = "Launch Conditions"
		else:
			caseStr = "On-Orbit"
		print("New case: {}% - {}".format(currentFillLevel,caseStr))
		upperHeading = document.add_heading("{}% Fill - {}".format(currentFillLevel,caseStr),2)
		document.add_picture("./TableSCs/{}.png".format(caseCounter),width=shared.Inches(6.3))
		document.add_page_break()
	lastFillLevel = currentFillLevel
	
	# print("{}% - {}g's".format(currentFillLevel,currentAcc))
	for j in range(1,data.shape[0]):
		currentTestName = data.loc[j].array[n]
		if isinstance(currentTestName, str):
			currentFrequency = data.loc[j].array[0]
			iterCounter += 1
			print("{}. {} - {} Hz - {} g's".format(iterCounter, currentTestName, currentFrequency, currentAcc))
			testNumber = int(currentTestName.split("test")[1])
			imagesPath = "./Plots/Report/{}/".format(currentTestName)
			heading = document.add_heading("Test {}: Fill Level = {}%, Acceleration = {} g's, Frequency = {:g} Hz".format(testNumber,currentFillLevel,currentAcc,currentFrequency), 3)
			for i in range(len(imagesOrder)):
				pictureName = "{}_{}".format(currentTestName,imagesOrder[i])
				document.add_picture('{}{}.png'.format(imagesPath,pictureName),width=shared.Inches(6.3),height=shared.Inches(4))
				run = document.add_paragraph("Zoom from t=15s to t=22s")
				document.add_picture('{}{}-zoomed.png'.format(imagesPath,pictureName),width=shared.Inches(6.3),height=shared.Inches(4))

sections = document.sections
for section in sections:
	section.top_margin = shared.Inches(1)
	section.bottom_margin = shared.Inches(1)
	section.left_margin = shared.Inches(1)
	section.right_margin = shared.Inches(1)
document.save(documentTitle)